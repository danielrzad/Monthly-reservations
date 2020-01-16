from docx.api import Document
from docx.shared import Pt
from pprint import pprint 
from datetime import datetime
import glob
import calendar


data = {}

# List of MS .docx files to iterate through
docx_files_list = []
for name in glob.glob('*.docx'):
    docx_files_list.append(name)

print(docx_files_list)

# Keys to data dictionary
keys = []
header_table = Document(docx_files_list[0]).tables[0]
for i, txt in enumerate(header_table.rows[0].cells):
    keys.append(txt.text)

# Creating data dictionary in the pattern like that
# data = {'Adamowicz1':{nazwisko: adamowicz}
#			'Adamowicz2':{nazwisko: adamowicz}
for file in docx_files_list:
	document = Document(file)
	table = document.tables[0]
	key = ''
	repetitions = 1
	for i, row in enumerate(table.rows[1:]):
	    text = [cell.text for cell in row.cells]
	    if text[1] == '':
	    	key = 'PUSTE'
	    if text[1] == key:
	    	repetitions += 1
	    else:
	    	key = text[1]
	    	repetitions = 1
	    name_dict = dict(zip(keys, text))
	    data[key + str(repetitions)] = name_dict

pprint(data)

def day_to_num(day_of_the_week):
	days = ['PONIEDZIAŁEK', 'WTOREK', 'ŚRODA', \
		'CZWARTEK', 'PIĄTEK', 'SOBOTA', 'NIEDZIELA']
	for c, v in enumerate(days):
		if v == day_of_the_week:
			return c

def get_days(year, month, day):
	day = day_to_num(day)
	# day is:[0-Monday,1-Tuesday,2-Wednesday,3-Thursday
	# 		 [4-Friday,5-Saturday,6-Sunday]
	c = calendar.Calendar()
	r =  [date for date, _day in c.itermonthdays2(year, month) \
		if date != 0 and _day == day]
	return str(r).strip('[]')

def add_new(dic_keys, week_day, hour, name):
	for i in range(1, 4):
		player_vals = ['', name, week_day, hour,
				 		get_days(datetime.now().year,
				 				datetime.now().month,
				 				week_day),
				 		'', '']
		player_data = dict(zip(dic_keys, player_vals))
		data[name + str(i)] = player_data

def days_update(dd):
	for i in dd.keys():
		dd[i]['DNI KALENDARZOWE'] = get_days(2020,
									1,
									dd[i]['DZIEŃ TYGODNIA'])

days_update(data)

def del_existing(dict_name, name):
	for i in range(1, 4):
		del dict_name[name + str(i)]


for key in data.keys():
	data[key.upper()] = data.pop(key)

del_existing(data, 'SUROWANIEC\n(ADAMOWICZ)')
del_existing(data, 'PRYCIAK')
del_existing(data, 'PYZARA')

data[key.upper()] = data.pop(key)

for i in range(1, 4):
	data['LUPICKI' + str(i)] = data.pop('ŁUPICKI' + str(i))
	data['TRACZ\n(BARTEK)\n(DZIECKO)' + str(i)] = data.pop('BARTEK\n(TRACZ)' + str(i))
	data['WITTWER\n(CIURA)' + str(i)] = data.pop('CIURA\n(WITTWER)' + str(i))
	data['TRACZ\n(KAMIL)' + str(i)] = data.pop('KAMIL\n(TRACZ)' + str(i))
	data['TRACZ\n(NATAN)' + str(i)] = data.pop('NATAN\n(TRACZ)' + str(i))
	data['WITTWER\n(MARIA)' + str(i)] = data.pop('MARIA\n(WITTWER) ' + str(i))
	data['TRACZ\n(PAWEŁ)' + str(i)] = data.pop('PAWEŁ\n(TRACZ)' + str(i))


add_new(keys, 'ŚRODA', '14.30-15.30', 'CIAK(PYZARA)')
add_new(keys, 'PONIEDZIAŁEK', '17.00-18.00', 'CIAK\n(PRYWATNE)')
add_new(keys, 'WTOREK', '8.00-9.00', 'TRACZ\n(JÓZEF)')
add_new(keys, 'ŚRODA', '14.30-15.30', 'TRACZ\n(WIKTOR)')
add_new(keys, 'CZWARTEK', '13.30-14.30', 'TRACZ\n(BARTEK)\n(DZIECKO)')
add_new(keys, 'ŚRODA', '13.30-14.30', 'TRACZ\n(KORDAS)')
add_new(keys, 'CZWARTEK', '12.00-13.00', 'WITTWER\n(KARAPETYAN)')


data_keys_sorted = sorted(data.keys())

data_len = len(data.keys())
docx_needed = data_len / 24

for c, v in enumerate(sorted(data.keys()), 1):
		data[v]['LP'] = c


def docx_fill(docx_lst, dd, sorted_keys, docx_num):
	# dd = data_dictionary
	key_order = ['LP', 'NAZWISKO', 'DZIEŃ TYGODNIA',
				'GODZINA', 'DNI KALENDARZOWE', 'NALEŻNOŚĆ\nZA GODZINĘ',
				'UWAGI', 'PŁATNOŚĆ']
	if docx_num.is_integer() != True:
		docx_num = int(docx_num) + 1
	current_docx = 0
	for c, v in enumerate(sorted_keys):
		if c in [8, 16, 24, 32, 40, 48, 56, 64]:
			current_docx += 1
		docx = Document(docx_lst[current_docx])
		tbl = docx.tables[0]
		for i, row in enumerate(tbl.rows[1:]):
			for j, cell in enumerate(row.cells):
				cell.text = str(dd[sorted_keys[c]][key_order[j]])
		document.save('REZERWACJE STAŁE' + str(current_docx) + '.docx')

docx_fill(docx_files_list, data, data_keys_sorted, docx_needed)
pprint(data)