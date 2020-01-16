import calendar
from docx.api import Document
from docx.shared import Pt
from pprint import pprint 
from os import listdir
from os.path import join

def main(path_to_old_docxs, save_path, year, month, nw_days):

    def prev_month_docxs(path_to_old_docxs):
        # searchs for prev month docx files, returning paths to them as a list
        path = path_to_old_docxs
        docx_files_list = [join(path, f) for f in listdir(path) if f.endswith(".docx")]
        return docx_files_list

    old_docx_paths = prev_month_docxs(path_to_old_docxs)

    def table_column_header_keys(sample_docx):
        # reads keys from first row of the table, which later will be keys in data-dict
        keys = []
        header_table = Document(sample_docx).tables[0]
        header_cells = header_table.row_cells(0)
        for i, cell in enumerate(header_cells):
            keys.append(cell.text)
        return keys

    inner_keys = table_column_header_keys(old_docx_paths[0])

    def data_dict(docxs_list, keys_list):
        # 1) reading data from MS docx files
        # 2) parsing them into dictionary 
        # 3) data_structure:
        # data['PLAYER NAME'] = {'INNER KEY 1': INNER VALUE 1,
        #                       'INNER KEY 2': INNER VALUE 2, etc.}
        data = {}
        for file in docxs_list:
            document = Document(file)
            table = document.tables[0]
            empty_keys = 0
            repetitions = 1
            key = '*'
            for i, row in enumerate(table.rows[1:]):
                text = [cell.text for cell in row.cells]
                if text[1] == '':
                    empty_keys += 1
                    key = 'ZZ' + str(empty_keys)
                elif text[1] == key[:-1]:
                    repetitions += 1
                    key = text[1] + str(repetitions)
                else:
                    repetitions = 1
                    key = text[1] + str(repetitions)
                name_dict = dict(zip(keys_list, text))
                data[key] = name_dict
        return data

    # data variable stores nested dictionary in which players data are stored
    data = data_dict(old_docx_paths, inner_keys)

    def calendar_days_fill(year, month, day, nw_days):
        # nw_days = not-working days which we don't want in the data
        # fills values at data dictionary inner key 'DNI KALENDARZOWE'

        def day_to_num(day):
            # changes input from data dictionary
            # basicly changes word to it's number representation
            days = ['PONIEDZIAŁEK', 'WTOREK', 'ŚRODA', \
                'CZWARTEK', 'PIĄTEK', 'SOBOTA', 'NIEDZIELA']
            for c, v in enumerate(days):
                if v == day:
                    return c

        def get_days(year, month, day, nw_days):
            # returning list of the calendar days based on given week day
            # excluding not working days when tennis courts are closed
            day = day_to_num(day)
            # day is:[0-Monday,1-Tuesday,2-Wednesday,3-Thursday
            #        [4-Friday,5-Saturday,6-Sunday]
            c = calendar.Calendar()
            r =  [date for date, _day in c.itermonthdays2(year, month) \
                if date != 0 and _day == day and date not in nw_days]
            return str(r).strip('[]')
        
        return get_days(year, month, day, nw_days)


    def add_new(data_dic, minor_data_keys, week_day, hour, name, year, month):
        # adding new player to data dictionary
        for i in range(1, 4):
            player_vals = ['', name, week_day, hour,
                            calendar_days_fill(year, month, week_day),
                            '', '']
            player_data = dict(zip(minor_data_keys, player_vals))
            data_dic[name + str(i)] = player_data
    
    def del_existing(data_dic, data_key):
        # deletes existing data key, key must be present in a data dictionary
        for i in range(1, 4):
            del data_dic[data_key + str(i)]

    def switch_data_key(data_dic, old_data_key, new_data_key):
        # changes old data key name to the new one 
        for i in range(1, 4):
            data_dic[new_data_key + str(i)] = data_dic.pop(old_data_key + str(i))

    def docx_fill(old_docxs, save_path, data_dict):
        sorted_data = sorted(data_dict.keys())
        sliced_data = [sorted_data[i * 24:(i + 1) * 24] for i in range((len(sorted_data) + 24 - 1) // 24)]
        key_order = ['LP', 'NAZWISKO', 'DZIEŃ TYGODNIA',
                     'GODZINA', 'DNI KALENDARZOWE', 'NALEŻNOŚĆ\nZA GODZINĘ',
                      'UWAGI', 'PŁATNOŚĆ']
        for slicee in sliced_data:
            first_letter = slicee[0][0]
            last_letter = slicee[-1][0]
            document = Document(old_docxs[0])
            sample_table = document.tables[0]
            padding_letters = set()
            for i, row in enumerate(sample_table.rows[1:]):
                padding_letters.add(data_dict[slicee[i]]['NAZWISKO'][:1])
                for j, cell in enumerate(row.cells):
                    cell.text = data_dict[slicee[i]][key_order[j]]
                    #cell.paragraphs[0].style = document.styles['Normal']
            paragraphs = document.paragraphs
            paragraphs[0].text = '  -  '.join(sorted(list(padding_letters)))
            document.save(save_path + 'REZERWACJE STAŁE  ' + first_letter + '-' + last_letter + '.docx')

    switch_data_key(data, 'ŁUPICKI', 'LUPICKI')
    
    for i in data.keys():
        data[i]['DNI KALENDARZOWE'] = calendar_days_fill(year, month, data[i]['DZIEŃ TYGODNIA'], [1, 6])

    docx_fill(old_docx_paths, save_path, data)


# def main(path_to_old_docxs, save_path, year, month, nw_days)
# year -- 4 digits like 2020, month -- 1-12 (1==Jan)
# nw_days -- not working days, on which tennis courts are closed
main('old/', 'C:\\Users\\Admin\\Desktop\\merging docx tables\\new\\', 2020, 1, [1, 6])